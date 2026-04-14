"""
Build PCC Calorie Tracker Word poster.
Matches the poster.html style: dark-green left panel, parchment right panel,
Cormorant Garamond headings, Lora body, amber bullet accents.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from lxml import etree

# ── Palette (from site CSS) ──────────────────────────────────────────────────
GREEN_DARK  = '1F4330'   # left panel bg
BARK        = '19120B'   # main text
EARTH       = '5A4A3A'   # body text on light bg
CLAY        = '9A8878'   # muted text
AMBER       = 'B8924A'   # bullet accent
FERN        = '2F5E42'   # italic accent in headings
WHITE_WARM  = 'FFFEF9'   # text on dark bg
CREAM_TEXT  = 'D0CABC'   # muted text on dark bg (approx rgba(255,253,245,0.73))
PARCHMENT   = 'F5F0E6'   # right panel bg

ASSETS = '/Users/micah/Documents/Calorie Tracker/public/assets'

# ── Helpers ──────────────────────────────────────────────────────────────────
def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    shd = parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        f' w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
    )
    tcPr.append(shd)


def cell_margins(cell, top_pt, right_pt, bottom_pt, left_pt):
    def twips(pt): return int(pt * 20)  # 1 pt = 20 twips
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcMar'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(parse_xml(
        f'<w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:top w:w="{twips(top_pt)}" w:type="dxa"/>'
        f'<w:right w:w="{twips(right_pt)}" w:type="dxa"/>'
        f'<w:bottom w:w="{twips(bottom_pt)}" w:type="dxa"/>'
        f'<w:left w:w="{twips(left_pt)}" w:type="dxa"/>'
        f'</w:tcMar>'
    ))


def para_spacing(para, before_pt=0, after_pt=0):
    pPr = para._p.get_or_add_pPr()
    existing = pPr.find(qn('w:spacing'))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(parse_xml(
        f'<w:spacing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        f' w:before="{int(before_pt*20)}" w:after="{int(after_pt*20)}" w:line="276" w:lineRule="auto"/>'
    ))


def styled_run(para, text, font_name, size_pt, hex_color,
               bold=False, italic=False):
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor.from_string(hex_color)
    run.font.bold = bold
    run.font.italic = italic
    return run


def add_para(cell, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=6):
    p = cell.add_paragraph()
    p.alignment = align
    para_spacing(p, before_pt, after_pt)
    return p


def kill_default_para(cell):
    """Remove the blank paragraph Word inserts into every new cell."""
    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)


def no_table_borders(table):
    tblPr = table._tbl.tblPr
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(parse_xml(
        '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    ))


def no_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(parse_xml(
        '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="none" w:sz="0"/>'
        '<w:left w:val="none" w:sz="0"/>'
        '<w:bottom w:val="none" w:sz="0"/>'
        '<w:right w:val="none" w:sz="0"/>'
        '</w:tcBorders>'
    ))


# ── Document setup ────────────────────────────────────────────────────────────
doc = Document()

# Wipe the default empty paragraph
for p in list(doc.paragraphs):
    p._element.getparent().remove(p._element)

section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(0)
section.right_margin  = Inches(0)
section.top_margin    = Inches(0)
section.bottom_margin = Inches(0)
section.header_distance = Pt(0)
section.footer_distance = Pt(0)

# ── Two-column table (full page) ──────────────────────────────────────────────
PAGE_W_EMU  = int(8.5 * 914400)
LEFT_W_EMU  = int(PAGE_W_EMU * 0.52)
RIGHT_W_EMU = PAGE_W_EMU - LEFT_W_EMU

table = doc.add_table(rows=1, cols=2)
no_table_borders(table)

# Fix table width
tblPr = table._tbl.tblPr
old_w = tblPr.find(qn('w:tblW'))
if old_w is not None:
    tblPr.remove(old_w)
tblPr.append(parse_xml(
    f'<w:tblW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
    f' w:w="{int(PAGE_W_EMU / 914400 * 1440)}" w:type="dxa"/>'
))

# Fix column widths
table.columns[0].width = LEFT_W_EMU
table.columns[1].width = RIGHT_W_EMU

left  = table.cell(0, 0)
right = table.cell(0, 1)

shade_cell(left,  GREEN_DARK)
shade_cell(right, PARCHMENT)
no_cell_borders(left)
no_cell_borders(right)

# Padding: 0.32in sides, 0.28in top/bottom
PAD_TB = 20   # pt
PAD_LR = 23   # pt
cell_margins(left,  PAD_TB, PAD_LR, PAD_TB, PAD_LR)
cell_margins(right, PAD_TB, PAD_LR, PAD_TB, PAD_LR)

kill_default_para(left)
kill_default_para(right)

# ── LEFT PANEL ────────────────────────────────────────────────────────────────

# App brand badge
brand = add_para(left, after_pt=10)
styled_run(brand, '⊛  PCC Calorie Tracker', 'Lora', 8.5, CREAM_TEXT)

# Main headline (multi-run for the italic portion)
headline = add_para(left, after_pt=10)
headline._p.get_or_add_pPr()
styled_run(headline, 'Track meals from ', 'Cormorant Garamond', 42, WHITE_WARM)
styled_run(headline, 'PCC dining menus', 'Cormorant Garamond', 42, CREAM_TEXT, italic=True)
styled_run(headline, '\nin seconds.', 'Cormorant Garamond', 42, WHITE_WARM)

# Thin divider rule (em-dash line, styled in amber)
rule = add_para(left, after_pt=8)
styled_run(rule, '─────────────', 'Lora', 8, AMBER)

# Sub-copy
sub = add_para(left, after_pt=12)
styled_run(sub, (
    'Built for PCC students. Pick items straight from your dining '
    'menu and your calories and macros are filled in automatically.'
), 'Lora', 9.5, CREAM_TEXT)

# Bullet points
BULLETS = [
    'Connect to live PCC menu items',
    'Log meals faster with pre-filled nutrition',
    'Keep goals on track with Health and photo logging',
]
for text in BULLETS:
    bp = add_para(left, after_pt=5)
    styled_run(bp, '●  ', 'Lora', 7.5, AMBER)
    styled_run(bp, text, 'Lora', 9.5, WHITE_WARM)

# Hero screenshot
img_para = add_para(left, WD_ALIGN_PARAGRAPH.LEFT, before_pt=14, after_pt=0)
img_para.add_run().add_picture(f'{ASSETS}/screenshots/IMG_5922.PNG', width=Inches(2.1))

# ── RIGHT PANEL ───────────────────────────────────────────────────────────────

# Spacer
sp = add_para(right, after_pt=4)
styled_run(sp, ' ', 'Lora', 6, CLAY)

# Kicker
kicker = add_para(right, WD_ALIGN_PARAGRAPH.CENTER, after_pt=4)
styled_run(kicker, 'PCC CALORIE TRACKER', 'Lora', 7, CLAY)
# Letter spacing via XML
rPr = kicker.runs[0]._r.get_or_add_rPr()
rPr.append(parse_xml(
    '<w:spacing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="100"/>'
))

# Pull-quote headline
pull = add_para(right, WD_ALIGN_PARAGRAPH.CENTER, after_pt=12)
styled_run(pull, 'PCC menu connection,\n', 'Cormorant Garamond', 28, BARK)
styled_run(pull, 'built in.', 'Cormorant Garamond', 28, FERN, italic=True)

# QR code
qr_para = add_para(right, WD_ALIGN_PARAGRAPH.CENTER, after_pt=8)
qr_para.add_run().add_picture(f'{ASSETS}/qr/website-qr-new.png', width=Inches(1.72))

# Scan CTA
scan = add_para(right, WD_ALIGN_PARAGRAPH.CENTER, after_pt=3)
styled_run(scan, 'Scan to open the app page', 'Lora', 10.5, BARK, bold=True)

# URL
url = add_para(right, WD_ALIGN_PARAGRAPH.CENTER, after_pt=14)
styled_run(url, 'calorietrackerpcc.com', 'Lora', 8.5, CLAY)

# Menu screenshot
menu_para = add_para(right, WD_ALIGN_PARAGRAPH.CENTER, before_pt=2, after_pt=8)
menu_para.add_run().add_picture(f'{ASSETS}/screenshots/menu-latest.png', width=Inches(2.1))

# Caption
caption = add_para(right, WD_ALIGN_PARAGRAPH.CENTER, after_pt=0)
styled_run(caption, 'Real menu view inside the app', 'Lora', 7.5, CLAY, italic=True)

# Footer note
footer = add_para(right, WD_ALIGN_PARAGRAPH.CENTER, before_pt=10, after_pt=0)
styled_run(footer, 'iOS app  ·  PCC Dining menus  ·  Calorie + nutrient tracking', 'Lora', 7.5, CLAY)

# ── Save ──────────────────────────────────────────────────────────────────────
out = '/Users/micah/Documents/Calorie Tracker/PCC Calorie Tracker Poster.docx'
doc.save(out)
print(f'Saved → {out}')
